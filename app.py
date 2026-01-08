# =========================================
# MULTI-SUBJECT REPORT COMMENT GENERATOR - Secure Streamlit Version
# Supports English and Science, Years 7 & 8
# Enhanced with security, privacy, and UX features
# =========================================

import random
import streamlit as st
from docx import Document
import tempfile
import os
import time
from datetime import datetime, timedelta
import pandas as pd
import zipfile
import io

# ========== SECURITY & PRIVACY SETTINGS ==========
TARGET_CHARS = 499  # target character count including spaces
MAX_FILE_SIZE_MB = 5  # Maximum upload size
MAX_ROWS_PER_UPLOAD = 100  # Maximum students per upload
RATE_LIMIT_SECONDS = 10  # Minimum time between uploads

# ========== PAGE CONFIGURATION ==========
st.set_page_config(
    page_title="üîí Secure Report Generator",
    page_icon="üìö",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========== SECURITY INITIALIZATION ==========
# Clear session state on fresh load
if 'app_initialized' not in st.session_state:
    st.session_state.clear()
    st.session_state.app_initialized = True
    st.session_state.upload_count = 0
    st.session_state.last_upload_time = datetime.now()
    st.session_state.generated_files = []

# ========== IMPORT STATEMENTS ==========
try:
    from statements_year7_English import (
        opening_phrases as opening_7_eng,
        attitude_bank as attitude_7_eng,
        reading_bank as reading_7_eng,
        writing_bank as writing_7_eng,
        reading_target_bank as target_7_eng,
        writing_target_bank as target_write_7_eng,
        closer_bank as closer_7_eng
    )

    from statements_year8_English import (
        opening_phrases as opening_8_eng,
        attitude_bank as attitude_8_eng,
        reading_bank as reading_8_eng,
        writing_bank as writing_8_eng,
        reading_target_bank as target_8_eng,
        writing_target_bank as target_write_8_eng,
        closer_bank as closer_8_eng
    )

    from statements_year7_science import (
        opening_phrases as opening_7_sci,
        attitude_bank as attitude_7_sci,
        science_bank as science_7_sci,
        target_bank as target_7_sci,
        closer_bank as closer_7_sci
    )

    from statements_year8_science import (
        opening_phrases as opening_8_sci,
        attitude_bank as attitude_8_sci,
        science_bank as science_8_sci,
        target_bank as target_8_sci,
        closer_bank as closer_8_sci
    )
except ImportError as e:
    st.error(f"Missing required statement files: {e}")
    st.stop()

# ========== SECURITY FUNCTIONS ==========
def validate_upload_rate():
    """Prevent rapid-fire uploads/abuse"""
    time_since_last = datetime.now() - st.session_state.last_upload_time
    if time_since_last < timedelta(seconds=RATE_LIMIT_SECONDS):
        wait_time = RATE_LIMIT_SECONDS - time_since_last.seconds
        st.error(f"Please wait {wait_time} seconds before uploading again")
        return False
    return True

def sanitize_input(text, max_length=100):
    """Sanitize user input to prevent injection attacks"""
    if not text:
        return ""
    # Remove special characters except basic punctuation and letters
    sanitized = ''.join(c for c in text if c.isalnum() or c in " .'-")
    # Truncate to max length
    return sanitized[:max_length].strip().title()

def validate_file(file):
    """Validate uploaded file size and type"""
    if file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
        return False, f"File too large (max {MAX_FILE_SIZE_MB}MB)"
    
    if not file.name.lower().endswith('.csv'):
        return False, "Only CSV files allowed"
    
    return True, ""

def process_csv_securely(uploaded_file):
    """Process CSV with auto-cleanup of temp files"""
    # Create temp file that auto-deletes
    with tempfile.NamedTemporaryFile(delete=False, suffix='.csv', mode='wb') as tmp:
        tmp.write(uploaded_file.getvalue())
        temp_path = tmp.name
    
    try:
        # Read CSV with limits
        df = pd.read_csv(temp_path, nrows=MAX_ROWS_PER_UPLOAD + 1)
        
        # Check row count
        if len(df) > MAX_ROWS_PER_UPLOAD:
            st.warning(f"Only processing first {MAX_ROWS_PER_UPLOAD} rows")
            df = df.head(MAX_ROWS_PER_UPLOAD)
        
        # Sanitize student names
        if 'Student Name' in df.columns:
            df['Student Name'] = df['Student Name'].apply(lambda x: sanitize_input(str(x)))
        
        return df
        
    except Exception as e:
        st.error(f"Error reading CSV: {e}")
        return None
        
    finally:
        # Always delete temp file
        try:
            os.unlink(temp_path)
        except:
            pass

# ========== HELPER FUNCTIONS ==========
def get_pronouns(gender):
    gender = gender.lower()
    if gender == "male":
        return "he", "his"
    elif gender == "female":
        return "she", "her"
    return "they", "their"

def lowercase_first(text):
    return text[0].lower() + text[1:] if text else ""

def truncate_comment(comment, target=TARGET_CHARS):
    if len(comment) <= target:
        return comment
    truncated = comment[:target].rstrip(" ,;.")
    if "." in truncated:
        truncated = truncated[:truncated.rfind(".")+1]
    return truncated

def fix_pronouns_in_text(text, pronoun, possessive):
    """Fix gender pronouns in statement text"""
    if not text:
        return text
    
    # Fix the specific spelling mistake you found
    text = text.replace('approacshed', 'approached')
    text = text.replace('tshe', 'the')
    
    # Replace all pronoun variations
    text = text.replace('he', pronoun).replace('He', pronoun.capitalize())
    text = text.replace('his', possessive).replace('His', possessive.capitalize())
    text = text.replace('him', pronoun).replace('Him', pronoun.capitalize())
    text = text.replace('himself', f"{pronoun}self")
    text = text.replace('herself', f"{pronoun}self")
    
    return text

def generate_comment(subject, year, name, gender, att, achieve, target, pronouns, attitude_target=None):
    p, p_poss = pronouns
    name = sanitize_input(name)

    if subject == "English":
    if year == 7:
        # ========== DEBUG CODE START ==========
        st.write("üîç **DEBUG INFO for Year 7 English:**")
        st.write(f"Student: {name}")
        st.write(f"Target band selected: {target}")
        
        # Get the raw text from your statement file
        raw_target_text = target_7_eng[target]
        st.write(f"Raw text from statement file: '{raw_target_text}'")
        
        # Check for "tshe" typo
        if "tshe" in raw_target_text:
            st.error("‚ùå PROBLEM FOUND: 'tshe' is in the statement file!")
        else:
            st.success("‚úÖ Good: No 'tshe' in statement file")
        # ========== DEBUG CODE END ==========
        
        opening = random.choice(opening_7_eng)
        # FIXED: CORRECT pronoun handling for attitude
        attitude_text = fix_pronouns_in_text(attitude_7_eng[att], p, p_poss)
        attitude_sentence = f"{opening} {name} {attitude_text}"
        if not attitude_sentence.endswith('.'):
            attitude_sentence += '.'
            
            # FIXED: CORRECT pronoun handling for reading - ADD PRONOUN!
            reading_text = fix_pronouns_in_text(reading_7_eng[achieve], p, p_poss)
            # Ensure it starts with pronoun
            if reading_text[0].islower():  # If starts with verb, add pronoun
                reading_text = f"{p} {reading_text}"
            reading_sentence = f"In reading, {reading_text}"
            if not reading_sentence.endswith('.'):
                reading_sentence += '.'
            
            # FIXED: CORRECT pronoun handling for writing - ADD PRONOUN!
            writing_text = fix_pronouns_in_text(writing_7_eng[achieve], p, p_poss)
            # Ensure it starts with pronoun
            if writing_text[0].islower():  # If starts with verb, add pronoun
                writing_text = f"{p} {writing_text}"
            writing_sentence = f"In writing, {writing_text}"
            if not writing_sentence.endswith('.'):
                writing_sentence += '.'
            
            reading_target_text = fix_pronouns_in_text(target_7_eng[target], p, p_poss)
            reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
            if not reading_target_sentence.endswith('.'):
                reading_target_sentence += '.'
            
            writing_target_text = fix_pronouns_in_text(target_write_7_eng[target], p, p_poss)
            writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
            if not writing_target_sentence.endswith('.'):
                writing_target_sentence += '.'
            
            closer_sentence = random.choice(closer_7_eng)
            
        else:  # Year 8
            opening = random.choice(opening_8_eng)
            attitude_text = fix_pronouns_in_text(attitude_8_eng[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            if not attitude_sentence.endswith('.'):
                attitude_sentence += '.'
            
            # FIXED: CORRECT pronoun handling for reading - ADD PRONOUN!
            reading_text = fix_pronouns_in_text(reading_8_eng[achieve], p, p_poss)
            # Ensure it starts with pronoun
            if reading_text[0].islower():  # If starts with verb, add pronoun
                reading_text = f"{p} {reading_text}"
            reading_sentence = f"In reading, {reading_text}"
            if not reading_sentence.endswith('.'):
                reading_sentence += '.'
            
            # FIXED: CORRECT pronoun handling for writing - ADD PRONOUN!
            writing_text = fix_pronouns_in_text(writing_8_eng[achieve], p, p_poss)
            # Ensure it starts with pronoun
            if writing_text[0].islower():  # If starts with verb, add pronoun
                writing_text = f"{p} {writing_text}"
            writing_sentence = f"In writing, {writing_text}"
            if not writing_sentence.endswith('.'):
                writing_sentence += '.'
            
            reading_target_text = fix_pronouns_in_text(target_8_eng[target], p, p_poss)
            reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
            if not reading_target_sentence.endswith('.'):
                reading_target_sentence += '.'
            
            writing_target_text = fix_pronouns_in_text(target_write_8_eng[target], p, p_poss)
            writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
            if not writing_target_sentence.endswith('.'):
                writing_target_sentence += '.'
            
            closer_sentence = random.choice(closer_8_eng)

    else:  # Science
        if year == 7:
            opening = random.choice(opening_7_sci)
            attitude_text = fix_pronouns_in_text(attitude_7_sci[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            if not attitude_sentence.endswith('.'):
                attitude_sentence += '.'
            
            # FIXED: CORRECT pronoun handling for science - ADD PRONOUN!
            science_text = fix_pronouns_in_text(science_7_sci[achieve], p, p_poss)
            # Ensure it starts with pronoun
            if science_text[0].islower():  # If starts with verb, add pronoun
                science_text = f"{p} {science_text}"
            reading_sentence = science_text
            if not reading_sentence.endswith('.'):
                reading_sentence += '.'
            
            target_text = fix_pronouns_in_text(target_7_sci[target], p, p_poss)
            reading_target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            if not reading_target_sentence.endswith('.'):
                reading_target_sentence += '.'
            
            writing_target_sentence = ""  # Not used for science
            closer_sentence = random.choice(closer_7_sci)
            writing_sentence = ""  # Not used for science
            
        else:  # Year 8
            opening = random.choice(opening_8_sci)
            attitude_text = fix_pronouns_in_text(attitude_8_sci[att], p, p_poss)
            attitude_sentence = f"{opening} {name} {attitude_text}"
            if not attitude_sentence.endswith('.'):
                attitude_sentence += '.'
            
            # FIXED: CORRECT pronoun handling for science - ADD PRONOUN!
            science_text = fix_pronouns_in_text(science_8_sci[achieve], p, p_poss)
            # Ensure it starts with pronoun
            if science_text[0].islower():  # If starts with verb, add pronoun
                science_text = f"{p} {science_text}"
            reading_sentence = science_text
            if not reading_sentence.endswith('.'):
                reading_sentence += '.'
            
            target_text = fix_pronouns_in_text(target_8_sci[target], p, p_poss)
            reading_target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
            if not reading_target_sentence.endswith('.'):
                reading_target_sentence += '.'
            
            writing_target_sentence = ""  # Not used for science
            closer_sentence = random.choice(closer_8_sci)
            writing_sentence = ""  # Not used for science

    # optional attitude target
    if attitude_target:
        attitude_target = sanitize_input(attitude_target)
        attitude_target_sentence = f" {lowercase_first(attitude_target)}"
        if not attitude_target_sentence.endswith('.'):
            attitude_target_sentence += '.'
    else:
        attitude_target_sentence = ""

    comment_parts = [
        attitude_sentence + attitude_target_sentence,
        reading_sentence,
        writing_sentence,
        reading_target_sentence,
        writing_target_sentence,
        closer_sentence
    ]

    comment = " ".join([c for c in comment_parts if c])  # remove empty strings
    
    # FIXED: Ensure comment ends with a period
    comment = comment.strip()
    if not comment.endswith('.'):
        comment += '.'
    
    comment = truncate_comment(comment, TARGET_CHARS)
    
    # Double-check period after truncation
    if not comment.endswith('.'):
        comment = comment.rstrip(' ,;') + '.'
    
    return comment

# ========== STREAMLIT APP LAYOUT ==========

# Sidebar for navigation and info
with st.sidebar:
    st.title("üìö Navigation")
    
    app_mode = st.radio(
        "Choose Mode",
        ["Single Student", "Batch Upload", "Privacy Info"]
    )
    
    st.markdown("---")
    st.markdown("### üîí Privacy Features")
    st.info("""
    - No data stored on servers
    - All processing in memory
    - Auto-deletion of temp files
    - Input sanitization
    - Rate limiting enabled
    """)
    
    if st.button("üîÑ Clear All Data", type="secondary", use_container_width=True):
        st.session_state.clear()
        st.session_state.app_initialized = True
        st.session_state.upload_count = 0
        st.session_state.last_upload_time = datetime.now()
        st.success("All data cleared!")
        st.rerun()
    
    st.markdown("---")
    st.caption("v2.3 ‚Ä¢ Teacher-Friendly Edition")

# Main content area with logo
col1, col2 = st.columns([1, 4])

with col1:
    try:
        st.image("logo.png", use_column_width=True)  # Logo will auto-fit column
    except:
        st.markdown("""
        <div style='text-align: center;'>
            <div style='font-size: 72px;'>üìö</div>
        </div>
        """, unsafe_allow_html=True)

with col2:
    st.title("Multi-Subject Report Comment Generator")
    st.caption("~499 characters per comment ‚Ä¢ Secure ‚Ä¢ No data retention")

# Privacy notice banner
st.warning("""
**PRIVACY NOTICE:** All data is processed in memory only. No files are stored on our servers. 
Close browser tab to completely erase all data. For use with anonymized student data only.
""", icon="üîí")

# ========== NEW: ONE-LINE PROGRESS TRACKER ==========
st.subheader("üéØ Three Easy Steps")

if 'progress' not in st.session_state:
    st.session_state.progress = 1

# Create a clean one-line progress tracker
progress_cols = st.columns(3)
with progress_cols[0]:
    st.markdown(f"""
    <div style='text-align: center; padding: 10px; background-color: {'#e6f3ff' if st.session_state.progress == 1 else '#f0f2f6'}; border-radius: 10px;'>
        <h3>{'‚úÖ' if st.session_state.progress > 1 else '1Ô∏è‚É£'}</h3>
        <h4>Select</h4>
        <p>Choose student details</p>
    </div>
    """, unsafe_allow_html=True)

with progress_cols[1]:
    st.markdown(f"""
    <div style='text-align: center; padding: 10px; background-color: {'#e6f3ff' if st.session_state.progress == 2 else '#f0f2f6'}; border-radius: 10px;'>
        <h3>{'‚úÖ' if st.session_state.progress > 2 else '2Ô∏è‚É£'}</h3>
        <h4>Generate</h4>
        <p>Create comments</p>
    </div>
    """, unsafe_allow_html=True)

with progress_cols[2]:
    st.markdown(f"""
    <div style='text-align: center; padding: 10px; background-color: {'#e6f3ff' if st.session_state.progress == 3 else '#f0f2f6'}; border-radius: 10px;'>
        <h3>{'‚úÖ' if st.session_state.progress > 3 else '3Ô∏è‚É£'}</h3>
        <h4>Download</h4>
        <p>Export reports</p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# ========== SINGLE STUDENT MODE ==========
if app_mode == "Single Student":
    st.subheader("üë§ Single Student Entry")
    
    with st.form("single_student_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            subject = st.selectbox("Subject", ["English", "Science"])
            year = st.selectbox("Year", [7, 8])
            name = st.text_input("Student Name", placeholder="Enter first name only")
            gender = st.selectbox("Gender", ["Male", "Female"])
        
        with col2:
            # Using dropdowns instead of sliders for faster teacher input
            att = st.selectbox("Attitude Band", 
                             options=[90,85,80,75,70,65,60,55,40],
                             index=3)  # Default to 75
            
            achieve = st.selectbox("Achievement Band",
                                 options=[90,85,80,75,70,65,60,55,40],
                                 index=3)  # Default to 75
            
            target = st.selectbox("Target Band",
                                options=[90,85,80,75,70,65,60,55,40],
                                index=3)  # Default to 75
            
            st.caption("üí° Use dropdowns for faster input. Tab key moves between fields.")
        
        attitude_target = st.text_area("Optional Attitude Next Steps",
                                     placeholder="E.g., continue to participate actively in class discussions...",
                                     height=60)
        
        col_submit = st.columns([3, 1])
        with col_submit[1]:
            submitted = st.form_submit_button("üöÄ Generate Comment", use_container_width=True)
    
    if submitted and name:
        if not validate_upload_rate():
            st.stop()
        
        name = sanitize_input(name)
        pronouns = get_pronouns(gender)
        
        with st.spinner("Generating comment..."):
            comment = generate_comment(subject, year, name, gender, att, achieve, 
                                     target, pronouns, attitude_target)
            char_count = len(comment)
        
        st.session_state.progress = 2
        
        # Display comment with stats
        st.subheader("üìù Generated Comment")
        st.text_area("", comment, height=200, key="comment_display")
        
        col_stats = st.columns(3)
        with col_stats[0]:
            st.metric("Character Count", f"{char_count}/{TARGET_CHARS}")
        with col_stats[1]:
            st.metric("Words", len(comment.split()))
        with col_stats[2]:
            if char_count < TARGET_CHARS - 50:
                st.success("‚úì Good length")
            else:
                st.warning("Near limit")
        
        # Store in session
        if 'all_comments' not in st.session_state:
            st.session_state.all_comments = []
        
        student_entry = {
            'name': name,
            'subject': subject,
            'year': year,
            'comment': comment,
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
        }
        st.session_state.all_comments.append(student_entry)
        
        # Add another button
        if st.button("‚ûï Add Another Student", type="secondary"):
            st.session_state.progress = 1
            st.rerun()

# ========== BATCH UPLOAD MODE ==========
elif app_mode == "Batch Upload":
    st.subheader("üìÅ Batch Upload (CSV)")
    
    st.info("""
    **CSV Format Required:**
    - Columns: Student Name, Gender, Subject, Year, Attitude, Achievement, Target
    - Gender: Male/Female
    - Subject: English/Science
    - Year: 7 or 8
    - Bands: 90,85,80,75,70,65,60,55,40
    """)
    
    # Add example CSV download
    example_csv = """Student Name,Gender,Subject,Year,Attitude,Achievement,Target
Aseel,Female,English,7,75,80,85
Mohamed,Male,Science,8,80,75,80
Sarah,Female,English,7,85,90,85"""
    
    st.download_button(
        label="üì• Download Example CSV",
        data=example_csv,
        file_name="example_students.csv",
        mime="text/csv"
    )
    
    uploaded_file = st.file_uploader("Choose CSV file", type=['csv'])
    
    if uploaded_file:
        if not validate_upload_rate():
            st.stop()
        
        is_valid, msg = validate_file(uploaded_file)
        if not is_valid:
            st.error(msg)
            st.stop()
        
        with st.spinner("Processing CSV securely..."):
            df = process_csv_securely(uploaded_file)
        
        if df is not None:
            st.success(f"Processed {len(df)} students successfully")
            
            # Preview
            with st.expander("üìã Preview Data (First 5 rows)"):
                st.dataframe(df.head())
            
            if st.button("üöÄ Generate All Comments", type="primary"):
                if 'all_comments' not in st.session_state:
                    st.session_state.all_comments = []
                
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for idx, row in df.iterrows():
                    # Update progress
                    progress = (idx + 1) / len(df)
                    progress_bar.progress(progress)
                    status_text.text(f"Processing {idx + 1}/{len(df)}: {row.get('Student Name', 'Student')}")
                    
                    try:
                        pronouns = get_pronouns(str(row.get('Gender', '')).lower())
                        comment = generate_comment(
                            subject=str(row.get('Subject', 'English')),
                            year=int(row.get('Year', 7)),
                            name=str(row.get('Student Name', '')),
                            gender=str(row.get('Gender', '')),
                            att=int(row.get('Attitude', 75)),
                            achieve=int(row.get('Achievement', 75)),
                            target=int(row.get('Target', 75)),
                            pronouns=pronouns
                        )
                        
                        student_entry = {
                            'name': sanitize_input(str(row.get('Student Name', ''))),
                            'subject': str(row.get('Subject', 'English')),
                            'year': int(row.get('Year', 7)),
                            'comment': comment,
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
                        }
                        st.session_state.all_comments.append(student_entry)
                        
                    except Exception as e:
                        st.error(f"Error processing row {idx + 1}: {e}")
                
                progress_bar.empty()
                status_text.empty()
                st.session_state.progress = 2
                st.success(f"Generated {len(df)} comments!")
                st.session_state.last_upload_time = datetime.now()

# ========== PRIVACY INFO MODE ==========
elif app_mode == "Privacy Info":
    st.subheader("üîê Privacy & Security Information")
    
    st.markdown("""
    ### How We Protect Student Data
    
    **Data Handling:**
    - All processing happens in your browser's memory
    - No student data is sent to or stored on our servers
    - Temporary files are created and immediately deleted
    - No database or persistent storage is used
    
    **Security Features:**
    1. **Input Sanitization** - Removes special characters from names
    2. **Rate Limiting** - Prevents abuse of the system
    3. **File Validation** - Checks file size and type
    4. **Auto-Cleanup** - Temporary files deleted after processing
    5. **Memory Clearing** - All data erased on browser close
    
    **Best Practices for Users:**
    - Use only first names or student IDs
    - Close browser tab when finished to clear all data
    - Download reports immediately after generation
    - For maximum privacy, use on school-managed devices
    
    **Compliance:**
    - Designed for use with anonymized data
    - Suitable for FERPA/GDPR compliant workflows
    - No third-party data sharing
    """)
    
    if st.button("üñ®Ô∏è Print Privacy Notice", type="secondary"):
        privacy_text = """
        MULTI-SUBJECT REPORT GENERATOR - PRIVACY NOTICE
        
        Data Processing: All student data is processed locally in memory only.
        No data is transmitted to external servers or stored permanently.
        
        Data Retention: All data is cleared when the browser tab is closed.
        
        Security: Input sanitization and validation prevents data injection.
        
        Usage: For use with anonymized student data only.
        """
        st.text_area("Privacy Notice for Records", privacy_text, height=300)

# ========== DOWNLOAD SECTION ==========
if 'all_comments' in st.session_state and st.session_state.all_comments:
    st.session_state.progress = 3
    st.markdown("---")
    st.subheader("üì• Download Reports")
    
    total_comments = len(st.session_state.all_comments)
    st.info(f"You have {total_comments} generated comment(s)")
    
    # Preview comments
    with st.expander(f"üëÅÔ∏è Preview All Comments ({total_comments})"):
        for idx, entry in enumerate(st.session_state.all_comments, 1):
            st.markdown(f"**{idx}. {entry['name']}** ({entry['subject']} Year {entry['year']})")
            st.write(entry['comment'])
            st.markdown("---")
    
    # Download options
    col_dl1, col_dl2, col_dl3 = st.columns(3)
    
    with col_dl1:
        if st.button("üìÑ Word Document", use_container_width=True):
            doc = Document()
            doc.add_heading('Report Comments', 0)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            doc.add_paragraph(f'Total Students: {total_comments}')
            doc.add_paragraph('')
            
            for entry in st.session_state.all_comments:
                doc.add_heading(f"{entry['name']} - {entry['subject']} Year {entry['year']}", level=2)
                doc.add_paragraph(entry['comment'])
                doc.add_paragraph('')
            
            # Save to bytes
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="‚¨áÔ∏è Download Word File",
                data=bio.getvalue(),
                file_name=f"report_comments_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    with col_dl2:
        if st.button("üìä CSV Export", use_container_width=True):
            # Create CSV
            csv_data = []
            for entry in st.session_state.all_comments:
                csv_data.append({
                    'Student Name': entry['name'],
                    'Subject': entry['subject'],
                    'Year': entry['year'],
                    'Comment': entry['comment'],
                    'Generated': entry['timestamp']
                })
            
            df_export = pd.DataFrame(csv_data)
            csv_bytes = df_export.to_csv(index=False).encode('utf-8')
            
            st.download_button(
                label="‚¨áÔ∏è Download CSV",
                data=csv_bytes,
                file_name=f"report_comments_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv",
                use_container_width=True
            )
    
    with col_dl3:
        if st.button("üóëÔ∏è Clear & Start Over", type="secondary", use_container_width=True):
            st.session_state.all_comments = []
            st.session_state.progress = 1
            st.success("All comments cleared! Ready for new entries.")
            st.rerun()

# ========== FOOTER ==========
st.markdown("---")
footer_cols = st.columns([2, 1])
with footer_cols[0]:
    st.caption("¬© Report Generator v2.3 | For educational use only")
with footer_cols[1]:
    if st.button("‚ÑπÔ∏è Quick Help", use_container_width=True):
        st.info("""
        **Quick Help:**
        1. **Select**: Choose student details
        2. **Generate**: Create comments
        3. **Download**: Export reports
        
        **Hotkeys:**
        - Tab: Move between fields
        - Enter: Submit form
        
        Need help? Contact support.
        """)
